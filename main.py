import docx
from appJar import gui


doc_path = ""  # Replace with your document path
interviewer_title = "Intervieweuse/Interviewer"
id_title = "ID/NOID"

def process_document(doc_path, interviewer_title, id_title):
    # Load the document
    doc = docx.Document(doc_path)

    i = 0
    for para in doc.paragraphs:
        if para.text.find("00:") == -1:
            i += 1
            print(para.text)
        else:
            print(para.text)
            break

    # Add the table with the required structure
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'


    # Helper function to set bold and font size
    def set_bold(run, size=11):
        run.font.bold = True
        run.font.size = docx.shared.Pt(size)


    def set_size(run, size=9):
        run.font.size = docx.shared.Pt(size)


    # Process the document paragraphs and add them to the table
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]

        # Remove empty paragraphs
        if para.text.strip() == "":
            p = para._element
            p.getparent().remove(p)
            continue

        if para.text.startswith(interviewer_title) or para.text.startswith(id_title):
            # Determine if it's Intervieweuse or ID 17 and format accordingly
            if para.text.startswith(interviewer_title):
                row = table.add_row()
                row.cells[0].merge(row.cells[1])  # Merge the cells to create a single column
                new_para = row.cells[0].paragraphs[0]
                parts = para.text.split(' ', 1)
                run = new_para.add_run(parts[0].strip() + ' ')
                set_bold(run)
                if len(parts) > 1:
                    run = new_para.add_run(parts[1].strip())
                    set_size(run)

                # when removing the element, the i becomes the next, so no need to increment
                para._element.getparent().remove(para._element)
                # Add subsequent paragraphs to the first column
                while i < len(doc.paragraphs):
                    next_para = doc.paragraphs[i]
                    if next_para.text.startswith(interviewer_title) or next_para.text.startswith(id_title):
                        break
                    if next_para.text.strip():
                        new_para.add_run("\n" + next_para.text.strip())
                    next_para._element.getparent().remove(next_para._element)
            elif para.text.startswith(id_title):
                row = table.add_row()
                new_para = row.cells[0].paragraphs[0]
                parts = para.text.split(' ', 2)
                run = new_para.add_run(parts[0] + ' ' + parts[1])
                set_bold(run)
                if len(parts) > 2:
                    run = new_para.add_run(parts[2])
                    set_size(run)

                # when removing the element, the i becomes the next, so no need to increment
                para._element.getparent().remove(para._element)
                # Add subsequent paragraphs to the first column
                while i < len(doc.paragraphs):
                    next_para = doc.paragraphs[i]
                    if next_para.text.startswith(interviewer_title) or next_para.text.startswith(id_title):
                        break
                    if next_para.text.strip():
                        new_para.add_run("\n" + next_para.text.strip())
                    next_para._element.getparent().remove(next_para._element)
                row.cells[1].text = ""

    output_path = doc_path.replace(".docx", "_edited.docx")
    doc.save(output_path)


def press():

    doc_path = app.entry("File Path")
    interviewer_title = app.entry("Interviewer Title")
    id_title = app.entry("ID Title")

    if not doc_path.endswith(".docx"):
        app.errorBox("Invalid File", "Please select a .docx file.")
        return
    else:
        #process the document
        process_document(doc_path, interviewer_title, id_title)
        app.setLabel("status", "Processing Complete.")

def clear_status(widget):
    app.setLabel("status", "")

with gui("Word Document Parser", "700x400") as app:
    app.label("Made by Marc Eid")

    app.entry("File Path", label=True, default="Choose File", kind="file", change=clear_status)
    app.entry("Interviewer Title", label=True, default=interviewer_title)
    app.entry("ID Title", label=True, default=id_title)

    app.label("status", "")
    app.buttons(["Submit", "Cancel"], [press, app.stop])





